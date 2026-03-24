from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "guides" / "Dept_Growth_Tracker_Guide.docx"


ACCENT = RGBColor(0x1F, 0x4E, 0x79)
TEXT = RGBColor(0x16, 0x30, 0x40)
MUTED = RGBColor(0x5C, 0x75, 0x86)
SOFT_FILL = "DCEFFC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(44)
    r = p.add_run("DATA CONSULTANT DIVISION")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Department Growth Tracker")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Operations Guide")
    r.italic = True
    r.font.size = Pt(16)
    r.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    r = p.add_run(
        "Reference guide for using the Department Growth Tracker to manage team competency ratings, "
        "track earned achievements, review employee progress, support onboarding, and maintain reusable growth records."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run(f"Updated {datetime.now().strftime('%B %d, %Y')}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT

    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.font.color.rgb = TEXT
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.color.rgb = TEXT
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.color.rgb = TEXT


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.color.rgb = TEXT


def add_two_col_table(doc, rows, widths=(2.15, 4.75)):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_shading(hdr[0], SOFT_FILL)
    set_cell_shading(hdr[1], SOFT_FILL)
    set_cell_text(hdr[0], "Area", bold=True, color=TEXT, size=10)
    set_cell_text(hdr[1], "Notes", bold=True, color=TEXT, size=10)
    hdr[0].width = Inches(widths[0])
    hdr[1].width = Inches(widths[1])
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True, color=TEXT)
        set_cell_text(cells[1], right, color=TEXT)
    doc.add_paragraph()


def add_three_col_table(doc, headers, rows, widths=(1.75, 2.15, 3.0)):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], SOFT_FILL)
        set_cell_text(table.rows[0].cells[idx], label, bold=True, color=TEXT, size=10)
        table.rows[0].cells[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, color=TEXT)
    doc.add_paragraph()


def build_document():
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    if "Body Text" not in styles:
        styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
    styles["Body Text"].font.name = "Aptos"
    styles["Body Text"].font.size = Pt(10.5)

    for name, size in (("Heading 1", 16), ("Heading 2", 12), ("Heading 3", 10)):
        styles[name].font.name = "Aptos"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True

    add_title_page(doc)

    add_heading(doc, "Contents", level=1)
    add_numbered(
        doc,
        [
            "Purpose and scope",
            "Operating model",
            "Navigation and page overview",
            "Manager workflows",
            "Individual workflows",
            "Achievements and XP tracking",
            "Import, export, and save behavior",
            "Administration notes",
            "Troubleshooting",
        ],
    )

    add_heading(doc, "Purpose And Scope", level=1)
    add_body(
        doc,
        "The Department Growth Tracker is a browser-based internal tool used to measure employee growth across the Data "
        "Consultant role. It combines team competency ratings, employee dashboards, earned achievement tracking, and "
        "resource links in one app so managers and individual contributors can use the same workspace for development conversations."
    )
    add_body(
        doc,
        "The current version also folds achievement tracking directly into the growth tracker structure. This means the app now "
        "captures both capability depth, through ratings, and evidence of applied work, through achievements and XP."
    )

    add_heading(doc, "Operating Model", level=1)
    add_bullets(
        doc,
        [
            "Competency ratings show skill depth across plan administration, technical work, process execution, communication, compliance, systems, and leadership.",
            "Achievements act as visible proof that growth is showing up in real work, not only in rating discussions.",
            "Manager mode focuses on team review, profile administration, onboarding oversight, and backup handling.",
            "Individual mode focuses on one employee dashboard, the competency library, achievements, and role resources.",
            "All app data is stored locally in the browser unless it is exported as JSON.",
        ],
    )

    add_heading(doc, "Navigation And Page Overview", level=1)
    add_three_col_table(
        doc,
        ("Page", "Who Uses It", "Purpose"),
        [
            ("Team Matrix", "Manager", "Compare employees across all competency rows and quickly identify strength, risk, and development patterns."),
            ("Gap Analysis", "Manager", "Highlight lower-rated areas that need follow-up or targeted coaching."),
            ("Team Profiles", "Manager", "Review employee-level profile summaries and competency rollups."),
            ("Manage Team", "Manager", "Add members, edit profiles, switch modes, and run import or export actions."),
            ("My Dashboard", "Individual", "Show progress ring, level status, XP meter, earned achievements, and strongest categories for the active employee."),
            ("Competency Library", "Individual", "Read each competency description, why it matters, and what development looks like by level."),
            ("Achievements", "Manager and Individual", "Track earned milestones, XP totals, trophy case cards, and category or tier filters."),
            ("Resources", "Manager and Individual", "Launch walkthroughs, linked tools, and supporting references."),
        ],
    )

    add_heading(doc, "Manager Workflows", level=1)
    add_heading(doc, "1. Review Team Capability", level=2)
    add_numbered(
        doc,
        [
            "Open Team Matrix to compare employee ratings side by side.",
            "Use the matrix colors and level labels to identify areas where a new hire needs support or where a tenured employee is ready for stretch work.",
            "Open Gap Analysis to focus on lower-rated competencies that need action.",
            "Use Team Profiles when a manager needs a more narrative profile-level summary before a one-on-one or calibration conversation.",
        ],
    )

    add_heading(doc, "2. Maintain Profiles", level=2)
    add_two_col_table(
        doc,
        [
            ("Add Member", "Creates a new employee profile with role, join date, notes, and avatar selection."),
            ("Profile Editing", "Allows managers to update identity details and retain a single source of truth for each employee record."),
            ("Mode Switch", "Toggles between manager and individual experiences so the same app can support different audiences."),
            ("Profile Switching", "Moves the active view to another employee without needing a separate login flow."),
        ],
    )

    add_heading(doc, "3. Use Achievements In Coaching", level=2)
    add_body(
        doc,
        "The Achievements page is designed as the proof layer for growth conversations. Managers can award achievements by clicking a locked card. "
        "Once awarded, the card clears from its greyed-out state, moves into the trophy case, and contributes XP toward the next level."
    )
    add_bullets(
        doc,
        [
            "Use onboarding achievements to mark early ramp milestones.",
            "Use data and tools achievements to recognize clean execution and growing technical confidence.",
            "Use team and legend achievements more sparingly because they signal broader trust or sustained contribution.",
        ],
    )

    add_heading(doc, "Individual Workflows", level=1)
    add_heading(doc, "My Dashboard", level=2)
    add_body(
        doc,
        "The dashboard gives the active employee a condensed progress view. It combines the current role, competency level context, progress ring, "
        "earned achievement count, and an XP meter that shows how far the employee is from the next achievement level."
    )
    add_bullets(
        doc,
        [
            "The hero area shows current level context and top-line progress.",
            "Achievement Progress summarizes current XP, current level title, and remaining XP to the next level.",
            "The earned achievement showcase surfaces proof of recent or important wins directly on the dashboard.",
        ],
    )

    add_heading(doc, "Competency Library", level=2)
    add_body(
        doc,
        "The library explains what each competency means, why it matters, and what level growth looks like. This page is especially useful during onboarding "
        "because it gives new team members a shared language for expectations before they are expected to perform every workflow independently."
    )

    add_heading(doc, "Achievements And XP Tracking", level=1)
    add_two_col_table(
        doc,
        [
            ("Greyed-Out Cards", "Locked achievements are intentionally desaturated so earned wins stand out more clearly."),
            ("Award Action", "Clicking a locked card awards it to the selected employee profile."),
            ("Remove Action", "Clicking an earned card allows managers to remove it if it was awarded by mistake."),
            ("Trophy Case", "Shows earned achievements near the top of the page so progress reads like a visible display shelf."),
            ("XP Meter", "Aggregates achievement XP and shows progress to the next level."),
            ("Tier Filters", "Rookie, Pro, and Expert filters help managers review milestone mix rather than a single long list."),
        ],
    )

    add_heading(doc, "Import, Export, And Save Behavior", level=1)
    add_three_col_table(
        doc,
        ("Function", "Location", "Behavior"),
        [
            ("Automatic Save", "App state", "Profile ratings, achievement state, and current selections are stored in local browser storage."),
            ("Export Profile", "Manage Team", "Downloads one employee profile as JSON so the record can be shared or updated elsewhere."),
            ("Export All", "Manage Team", "Creates a JSON backup of the full team state, including profiles and app mode."),
            ("Import Backup", "Manage Team", "Loads a saved JSON file back into the app. Full backup imports replace current local state after confirmation."),
            ("Profile Merge Import", "Manage Team", "Single-profile imports merge ratings and earned achievements into an existing employee when IDs match."),
        ],
    )

    add_heading(doc, "Administration Notes", level=1)
    add_bullets(
        doc,
        [
            "The app ships with seeded demo data when no saved state exists, which helps with first-use orientation.",
            "Achievement data is stored at the profile level so each employee can keep independent XP and earned milestone history.",
            "Resource links are intended to support onboarding, reference use, and direct launches into related tools.",
            "The current build is a standalone HTML app, so distribution is simple and does not require a backend deployment path.",
        ],
    )

    add_heading(doc, "Troubleshooting", level=1)
    add_two_col_table(
        doc,
        [
            ("No Profiles Are Showing", "Open Manage Team and add a member, or import a saved backup JSON file."),
            ("Dashboard Looks Empty", "Confirm an active profile is selected. The individual dashboard depends on the active employee record."),
            ("Achievements Did Not Carry Over", "Use a current backup or profile import. Achievement history is stored in each profile's earned ID list."),
            ("Unexpected Old Data", "Clear the browser's local storage for this app if a completely fresh reset is needed, then reload."),
            ("Guide Or Launcher Link Looks Wrong", "Use the main workspace launcher index page and confirm it points to the Personal Active dept-growth-tracker path."),
        ],
    )

    add_body(
        doc,
        f"Guide generated automatically on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}."
    )

    doc.save(OUTPUT)
    print(f"Created guide: {OUTPUT}")


if __name__ == "__main__":
    build_document()
