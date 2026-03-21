from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "guides" / "Fixed_Width_Report_Extraction_Console_Guide.docx"


ACCENT = RGBColor(0x2D, 0x78, 0xA5)
TEXT = RGBColor(0x16, 0x30, 0x40)
MUTED = RGBColor(0x5C, 0x75, 0x86)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(44)
    r = p.add_run("DATA CONSULTANT DIVISION")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Fixed-Width Report Extraction Console")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Walkthrough Guide")
    r.italic = True
    r.font.size = Pt(16)
    r.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    r = p.add_run(
        "Technical reference and day-to-day operating guide for loading fixed-width provider "
        "files, validating report codes, reviewing parsed output, and exporting workbook results."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run(f"Updated {datetime.now().strftime('%B %d, %Y')}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT

    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.font.color.rgb = TEXT
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.color.rgb = TEXT
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.color.rgb = TEXT


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.color.rgb = TEXT


def add_two_col_table(doc, rows, widths=(2.1, 4.8)):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_shading(hdr[0], "DCEFFC")
    set_cell_shading(hdr[1], "DCEFFC")
    set_cell_text(hdr[0], "Area", bold=True, color=TEXT, size=10)
    set_cell_text(hdr[1], "Notes", bold=True, color=TEXT, size=10)
    hdr[0].width = Inches(widths[0])
    hdr[1].width = Inches(widths[1])
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True, color=TEXT)
        set_cell_text(cells[1], right, color=TEXT)
    doc.add_paragraph()


def add_three_col_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "DCEFFC")
        set_cell_text(table.rows[0].cells[idx], label, bold=True, color=TEXT, size=10)
        table.rows[0].cells[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, color=TEXT)
    doc.add_paragraph()


def build_document():
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    if "Body Text" not in styles:
        styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
    styles["Body Text"].font.name = "Aptos"
    styles["Body Text"].font.size = Pt(10.5)

    for name, size in (("Heading 1", 16), ("Heading 2", 12), ("Heading 3", 10)):
        styles[name].font.name = "Aptos"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True

    add_title_page(doc)

    add_heading(doc, "Contents", level=1)
    add_numbered(
        doc,
        [
            "Purpose and scope",
            "Workflow overview",
            "Screen-by-screen walkthrough",
            "Validation and parsing behavior",
            "Export behavior",
            "Administrative notes",
            "Troubleshooting",
        ],
    )

    add_heading(doc, "Purpose And Scope", level=1)
    add_body(
        doc,
        "The Fixed-Width Report Extraction Console is a standalone internal HTML utility used to read inbound "
        "provider flat files, classify lines by report code, parse those lines into structured columns, and export "
        "the results into a workbook. The tool is designed to replace manual record splitting and reduce the amount "
        "of cleanup work that would otherwise happen in Excel."
    )
    add_body(
        doc,
        "Everything runs locally in the browser. There is no server dependency for the main parsing workflow. Users "
        "select a configured data source, supply the case and plan identifiers used in output naming, review parse "
        "results, and export the workbook when the output looks correct."
    )

    add_heading(doc, "Workflow Overview", level=1)
    add_numbered(
        doc,
        [
            "Select the correct provider or source from the Data Source dropdown.",
            "Enter the case number and plan name that should appear in the export file name.",
            "Load either a single raw text file or, for Voya mode, the extract folder that contains EXTRACT_XX_123456.TXT files.",
            "Allow the console to validate the incoming lines and sort accepted rows into configured reports.",
            "Review accepted counts, discarded counts, report codes found, summary totals, and the parsed preview tables.",
            "If the discard panel shows a valid unexpected record that should be retained, add it to the custom report list and parse again.",
            "Export the workbook once the output looks correct and the case number passes validation.",
        ],
    )

    add_heading(doc, "Walkthrough Checkpoints", level=2)
    add_bullets(
        doc,
        [
            "Default source on launch is EMPOWER.",
            "File mode is used for standard single-file parsing. Folder mode is used for VOYA extract packages.",
            "The parsed report area shows previews only, so very large files remain readable in-browser.",
            "Workbook export is disabled when no reports are found, when the case number is invalid, or when the selected source does not support export.",
            "The FMC export and Source Mapping Export controls are visible but currently marked as coming soon.",
        ],
    )

    add_heading(doc, "Screen-By-Screen Walkthrough", level=1)
    add_heading(doc, "1. Header And Setup Controls", level=2)
    add_two_col_table(
        doc,
        [
            ("Data Source", "Chooses the active parsing rule set. The current app supports the embedded source configurations, including file-based providers and a VOYA folder mode."),
            ("Case Number", "Validated before export. The app formats case values on blur and blocks export if the pattern does not match an approved prefix structure."),
            ("Plan Name", "Used in the workbook filename and title area."),
            ("Select Input", "Loads a single fixed-width text file when the current source uses file mode."),
            ("Select Extract Folder", "Used for VOYA mode. The app scans supported extract filenames inside the chosen folder."),
        ],
    )

    add_heading(doc, "2. Report List Panel", level=2)
    add_body(
        doc,
        "The report list panel functions as the first sanity check after a parse. Each configured report code appears as a small card with the number of accepted rows found for that report. "
        "If a custom record was added from the discard panel, the report list marks it as Custom so the user can tell which entries were introduced during the current session."
    )

    add_heading(doc, "3. Summary Tiles And Discard Review", level=2)
    add_two_col_table(
        doc,
        [
            ("Total Lines", "Counts nonblank lines reviewed during the parse run."),
            ("Accepted Lines", "Counts lines that matched the selected source's expected reports and had a matching layout definition."),
            ("Discarded Lines", "Counts lines rejected because the record code was unexpected, blank, unsupported, or missing a layout definition."),
            ("Reports Found", "Lists only the report codes that actually produced accepted rows in the current run."),
            ("Discard Reasons", "Explains why rows were skipped. Unexpected records can be added directly back into the session using the Add To Report List action."),
        ],
    )

    add_heading(doc, "4. Summary Totals", level=2)
    add_body(
        doc,
        "For sources that support balance breakdowns, the lower summary area calculates participant count, grand total, source totals, and fund totals from the parsed rows. "
        "When a source does not support breakdown logic, the app intentionally shows the section as unavailable so the user does not misread missing totals as zero values."
    )

    add_heading(doc, "5. Parsed Reports Area", level=2)
    add_body(
        doc,
        "Each accepted report renders as its own preview block. The block header shows the report code and friendly report name when one exists. "
        "The preview table is intended for validation, not as the final deliverable; export should still be used for the actual workbook output."
    )
    add_bullets(
        doc,
        [
            "Use the preview to confirm column alignment and spot obvious source-rule mismatches.",
            "Expand or collapse individual report blocks to focus on the areas that matter.",
            "Preview content is especially useful after adding a custom report from the discard panel.",
        ],
    )

    add_heading(doc, "6. Export Rail", level=2)
    add_three_col_table(
        doc,
        headers=("Control", "Current State", "What It Means"),
        widths=(1.9, 1.5, 3.7),
        rows=[
            ("Export Workbook", "Active when valid", "Creates the main Excel workbook when reports exist, export is supported for the chosen source, and the case number is valid."),
            ("Upload Mapping File", "Available", "Loads a mapping file name into session state for downstream export flows."),
            ("FMC Export", "Placeholder", "Visible but still marked coming soon. It remains locked until a mapping file is loaded."),
            ("Source Mapping Export", "Placeholder", "Visible but currently disabled with a coming soon tooltip."),
        ],
    )

    add_heading(doc, "Validation And Parsing Behavior", level=1)
    add_body(
        doc,
        "The parser uses source-specific record resolution rules. Most file-mode sources identify the record code from a configured character range, while some sources use custom resolver logic "
        "to distinguish overlapping records. VOYA folder mode uses the extract code in the filename and can also apply report-specific record rules when individual layouts require them."
    )

    add_heading(doc, "Case Number Rules", level=2)
    add_three_col_table(
        doc,
        headers=("Prefix Group", "Expected Shape", "Operational Note"),
        widths=(1.7, 2.2, 3.2),
        rows=[
            ("QK, NQ", "5 digits + 3 spaces + 5 digits", "Export is blocked until the formatted value matches this structure."),
            ("TA, TE, TG, JK, TO", "6 digits + 2 spaces + 5 digits", "The field auto-normalizes to uppercase and applies final spacing on blur."),
            ("Any invalid value", "Rejected", "The case message turns into an error and export stays disabled."),
        ],
    )

    add_heading(doc, "Discard Handling", level=2)
    add_bullets(
        doc,
        [
            "Unexpected record: the line had a record code that is not currently in the expected report list for the selected source.",
            "Missing layout: the line matched a report code, but no field map exists for that record.",
            "Skipped file: used in folder mode when a file does not match the supported extract naming convention.",
            "Unsupported extract code: used in folder mode when the filename code does not map to a known VOYA layout.",
        ],
    )
    add_body(
        doc,
        "Unexpected records can be promoted into the live session with Add To Report List. This is useful when a source includes a valid record the current report list did not expect. "
        "Custom reports remain session-scoped unless the user clears them."
    )

    add_heading(doc, "Export Behavior", level=1)
    add_body(
        doc,
        "The main workbook export packages the parsed result into a downloadable XLSX file named with the formatted case number, plan name, and the Full Parse suffix. "
        "The workbook includes a title sheet plus one sheet per parsed report that survived validation."
    )
    add_bullets(
        doc,
        [
            "Export is disabled when no accepted reports were found.",
            "Export is disabled for sources that explicitly mark supportsExport as false, including VOYA folder mode.",
            "The browser download name is sanitized so blank or unsafe values do not break the file name.",
        ],
    )

    add_heading(doc, "Administrative Notes", level=1)
    add_bullets(
        doc,
        [
            "Source configuration is embedded directly in the standalone HTML app.",
            "FIDELITY uses custom record-resolution logic because several records share overlapping prefixes.",
            "VOYA configuration is assembled from the embedded VOYA layout library and operates in folder mode rather than single-file mode.",
            "Breakdown totals are source-aware and can be intentionally disabled for providers where the balance-summary logic is not active.",
            "Mapping import state is session-only and currently supports future FMC-oriented workflows rather than the main workbook export path.",
        ],
    )

    add_heading(doc, "Troubleshooting", level=1)
    add_three_col_table(
        doc,
        headers=("Symptom", "Likely Cause", "Recommended Action"),
        widths=(2.1, 2.4, 2.6),
        rows=[
            ("No matching reports found", "Wrong source selected or file layout does not match the active configuration", "Confirm the provider, reload the file, and compare the discard list against the expected report set."),
            ("High discarded count", "Unexpected records or unsupported lines are present", "Review discard reasons first. Add a custom report only if the record is actually valid for this workflow."),
            ("Export button stays disabled", "No accepted reports, invalid case number, or export-disabled source", "Confirm reports were found, correct the case number, and verify the source supports workbook export."),
            ("Folder mode does nothing", "Selected source is not VOYA or the folder files do not match the expected naming pattern", "Switch to VOYA and verify the files use EXTRACT_XX_123456.TXT naming."),
            ("Summary totals unavailable", "The selected source does not support breakdown calculations", "Treat this as expected behavior unless the source should have breakdown support configured."),
        ],
    )

    add_heading(doc, "Closeout Reminder", level=1)
    add_body(
        doc,
        "Before sending workbook output downstream, confirm the source selection, skim the report preview blocks, validate the discard panel, and verify that the generated filename uses the correct case number and plan name. "
        "That short final check catches most avoidable export mistakes."
    )

    return doc


def main():
    doc = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote guide to {OUTPUT}")


if __name__ == "__main__":
    main()
