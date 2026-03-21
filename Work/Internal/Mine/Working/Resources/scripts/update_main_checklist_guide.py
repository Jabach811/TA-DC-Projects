from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
GUIDE_PATH = ROOT / "guides" / "Main_Checklist_Generator_How_To_Guide.docx"
MARKER = "March 2026 Feature Update"


ACCENT = RGBColor(0x1F, 0x4E, 0x79)
TEXT = RGBColor(0x16, 0x30, 0x40)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def main():
    doc = Document(GUIDE_PATH)

    if any(p.text.strip() == MARKER for p in doc.paragraphs):
        print("Guide already contains the March 2026 feature update.")
        return

    doc.add_page_break()

    p = doc.add_paragraph()
    p.style = "Heading 1"
    r = p.add_run(MARKER)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "The Main Checklist Generator now includes several workflow helpers that make live plan execution "
        "less scattered and easier to manage during active implementation work."
    )
    r.font.size = Pt(10.5)
    r.font.color.rgb = TEXT

    bullets = [
        "Pinned task notes let users promote important task-level notes into a dedicated pinboard so reminders do not get buried inside a single phase.",
        "Freeform pins make it possible to store quick working notes that are not tied to a single checklist task.",
        "Pinned notes can be reordered directly in the pinboard, which is useful for triaging the most urgent follow-up items.",
        "Custom calendar items let users add their own milestones and reminders alongside generated checklist due dates.",
        "Custom calendar dates surface in both the mini calendar and the full calendar view, so user-added events stay visible throughout the app.",
        "The intake team directory was refreshed so TM and TC dropdowns use the newer contact list and email autofill values.",
    ]
    for text in bullets:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.style = "Heading 2"
    r = p.add_run("New Function Reference")
    r.font.color.rgb = ACCENT

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ("Function", "Where It Lives", "Why It Matters")
    for idx, label in enumerate(headers):
        hdr[idx].text = label
        shade(hdr[idx], "D9EAF7")
        run = hdr[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = TEXT
        run.font.size = Pt(10)

    rows = [
        (
            "Task Note Pinning",
            "Checklist task note area",
            "Turns a task note into a reusable pin so it remains visible while the rest of the checklist moves forward.",
        ),
        (
            "Pinboard",
            "Checklist workspace center rail",
            "Keeps important reminders, blockers, and ad hoc working notes in one dedicated place.",
        ),
        (
            "Freeform Pins",
            "Pinboard add form",
            "Lets users track reminders that are not tied to any one checklist task.",
        ),
        (
            "Custom Calendar Items",
            "Mini calendar and full calendar",
            "Adds manual milestones such as meetings, client checkpoints, and internal deadlines alongside generated plan dates.",
        ),
        (
            "Updated TM/TC Directory",
            "Intake screen team section",
            "Improves contact autofill accuracy when starting a new implementation plan.",
        ),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            run = cells[idx].paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.style = "Heading 2"
    r = p.add_run("Recommended Usage Pattern")
    r.font.color.rgb = ACCENT

    numbered = [
        "Generate or load the plan as usual from the intake screen.",
        "Use task notes for detail that belongs to a specific task, then pin only the items that need constant visibility.",
        "Create freeform pins for cross-phase reminders such as client follow-up, payroll confirmations, or internal escalation notes.",
        "Add custom calendar items for non-generated milestones like meetings, deadlines, or manual checkpoints.",
        "Use the dashboard, vault, and calendar views as rollups, but keep the pinboard as the day-to-day execution scratchpad.",
    ]
    for text in numbered:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT

    p = doc.add_paragraph()
    r = p.add_run(f"Feature update appended {datetime.now().strftime('%B %d, %Y')}.")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT

    doc.save(GUIDE_PATH)
    print(f"Updated guide: {GUIDE_PATH}")


if __name__ == "__main__":
    main()
